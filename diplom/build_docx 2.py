#!/usr/bin/env python3
"""Сборка ВКР в .docx из Markdown-файлов diplom/ и данных result.txt."""

from __future__ import annotations

import os
import re
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Pt, Cm

ROOT = Path(__file__).resolve().parent.parent
DIPLOM = ROOT / "diplom"
RESULT_TXT = ROOT / "kursovaia" / "kursovaia" / "result.txt"
OUT = DIPLOM / "VKR_Symbolicheskaya_regressiya_chernovik.docx"


def ensure_fig_4_2_png(result_path: Path) -> None:
    script = DIPLOM / "plot_fig_4_2.py"
    if script.exists() and result_path.exists():
        mpl_dir = DIPLOM / ".mplconfig"
        mpl_dir.mkdir(parents=True, exist_ok=True)
        env = {**os.environ, "MPLCONFIGDIR": str(mpl_dir)}
        subprocess.run([sys.executable, str(script), str(result_path)], check=False, env=env)


def add_centered_picture(doc: Document, image_path: Path, width_cm: float = 15.5) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(6)
    p.add_run().add_picture(str(image_path), width=Cm(width_cm))


def add_figure_caption(doc: Document, text: str) -> None:
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.first_line_indent = Cm(0)
    cap.paragraph_format.space_after = Pt(12)
    run = cap.add_run(strip_md_inline(text))
    run.italic = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), "Times New Roman")


def set_doc_defaults(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(14)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    pf.space_after = Pt(0)
    # Шрифт для кириллицы в Word
    r = normal._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), "Times New Roman")


def strip_md_inline(s: str) -> str:
    s = re.sub(r"\*\*(.+?)\*\*", r"\1", s)
    s = re.sub(r"`(.+?)`", r"\1", s)
    s = s.replace("\\#", "#")
    return s.strip()


def add_paragraph(doc: Document, text: str, style: str | None = None) -> None:
    text = strip_md_inline(text) if text else ""
    if not text:
        return
    p = doc.add_paragraph(text, style=style)
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_heading(doc: Document, text: str, level: int) -> None:
    text = strip_md_inline(text)
    doc.add_heading(text, level=level)


def process_markdown_file(doc: Document, path: Path, *, skip_title_comment: bool = False) -> None:
    raw = path.read_text(encoding="utf-8")
    lines = raw.splitlines()
    i = 0
    in_code = False
    code_buf: list[str] = []

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            if in_code:
                if code_buf:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Cm(0.5)
                    run = p.add_run("\n".join(code_buf))
                    run.font.name = "Courier New"
                    run.font.size = Pt(11)
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        if skip_title_comment and line.startswith("# Титульный лист"):
            i += 1
            continue

        # Реферат: без строки объёма (страницы в конце по заданию)
        if path.name == "01_titul_referat.md" and "Объём:" in line:
            i += 1
            continue

        if line.startswith("# "):
            add_heading(doc, line[2:].strip(), 1)
        elif line.startswith("## "):
            add_heading(doc, line[3:].strip(), 2)
        elif line.startswith("### "):
            add_heading(doc, line[4:].strip(), 3)
        elif line.strip() == "---":
            pass
        elif (m_img := re.match(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$", line.strip())):
            alt, rel = m_img.group(1), m_img.group(2)
            img_path = (path.parent / rel).resolve()
            if img_path.exists():
                add_centered_picture(doc, img_path)
                add_figure_caption(doc, alt)
            else:
                add_paragraph(doc, f"[Не найден файл рисунка: {rel} — выполните diplom/plot_fig_4_2.py.]")
        elif line.strip().startswith("|") and line.count("|") >= 2:
            rows: list[list[str]] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                row = [c.strip() for c in lines[i].split("|")[1:-1]]
                if not all(re.match(r"^-+$", c) for c in row):
                    rows.append(row)
                i += 1
            i -= 1
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = "Table Grid"
                for ri, row in enumerate(rows):
                    for ci, cell in enumerate(row):
                        table.rows[ri].cells[ci].text = strip_md_inline(cell)
        elif not line.strip():
            pass
        else:
            add_paragraph(doc, line)
        i += 1


def add_title_page_block(doc: Document, path: Path) -> None:
    """Титул: без служебного заголовка про 'перенос в Word'."""
    raw = path.read_text(encoding="utf-8")
    # До РЕФЕРАТ
    part = raw.split("# РЕФЕРАТ")[0]
    for block in part.split("\n"):
        line = block.strip()
        if not line or line.startswith("# Титульный"):
            continue
        p = doc.add_paragraph()
        run = p.add_run(strip_md_inline(block))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
    doc.add_page_break()


def add_referat(doc: Document, path: Path) -> None:
    raw = path.read_text(encoding="utf-8")
    part = "# РЕФЕРАТ" + raw.split("# РЕФЕРАТ", 1)[1]
    # Временный файл-логика через строки
    lines = part.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.strip() == "---":
            i += 1
            continue
        if "Объём:" in line:
            i += 1
            continue
        if line.startswith("# "):
            add_heading(doc, line[2:].strip(), 1)
        elif line.strip():
            add_paragraph(doc, line)
        i += 1


def parse_benchmark_row_line(line: str) -> tuple[str, str, str, str, str] | None:
    """Разбор строки сводки; десятичный разделитель в протоколе — запятая."""
    s = line.strip()
    if ": RMSE=" not in s or ", rel=" not in s or ", threshold=" not in s or "status=" not in s:
        return None
    try:
        name, rest = s.split(": RMSE=", 1)
        rmse, rest = rest.split(", rel=", 1)
        rel, rest = rest.split(", threshold=", 1)
        threshold, rest = rest.split(", rel≤", 1)
        status = rest.split("status=", 1)[1].strip()
        return (name.strip(), rmse.strip(), rel.strip(), threshold.strip(), status)
    except ValueError:
        return None


def parse_benchmark_rows(result_path: Path) -> list[tuple[str, str, str, str, str]]:
    if not result_path.exists():
        return []
    rows: list[tuple[str, str, str, str, str]] = []
    for line in result_path.read_text(encoding="utf-8").splitlines():
        parsed = parse_benchmark_row_line(line)
        if parsed:
            rows.append(parsed)
    return rows


def parse_formulas_summary(result_path: Path) -> list[tuple[str, str]]:
    """Короткий список: имя файла, строка формулы (если есть)."""
    if not result_path.exists():
        return []
    blocks = result_path.read_text(encoding="utf-8").split("Файл: ")
    out: list[tuple[str, str]] = []
    for b in blocks[1:]:
        first_line = b.split("\n", 1)[0].strip()
        name = first_line
        formula = ""
        for line in b.splitlines():
            if line.startswith("Найденная формула"):
                formula = line.split("]", 1)[-1].strip() if "]" in line else line
                break
        out.append((name, formula))
    return out


def add_experiment_section(doc: Document, result_path: Path) -> None:
    doc.add_page_break()
    add_heading(doc, "4.3 Результаты контрольного эксперимента по данным программного протокола", 2)
    add_paragraph(
        doc,
        "Ниже приведены сводные данные по результатам прогона программного комплекса "
        f"(файл протокола: «{result_path.name}»). Таблица 4.2 формируется автоматически при сборке документа.",
    )
    add_heading(doc, "Таблица 4.2 – Сводка контрольных порогов (benchmark summary)", 3)
    bench = parse_benchmark_rows(result_path)
    if not bench:
        add_paragraph(doc, "(Данные протокола не найдены — укажите путь к result.txt.)")
        return
    table = doc.add_table(rows=1 + len(bench), cols=5)
    table.style = "Table Grid"
    hdr = ["Файл .xlsx", "RMSE", "Отн. ошибка", "Порог RMSE", "Статус"]
    for j, h in enumerate(hdr):
        table.rows[0].cells[j].text = h
    for i, row in enumerate(bench, start=1):
        for j, val in enumerate(row):
            table.rows[i].cells[j].text = val

    add_paragraph(doc, "")
    add_heading(doc, "Таблица 4.3 – Найденные модели по контрольным файлам", 3)
    forms = parse_formulas_summary(result_path)
    table2 = doc.add_table(rows=1 + len(forms), cols=2)
    table2.style = "Table Grid"
    table2.rows[0].cells[0].text = "Файл"
    table2.rows[0].cells[1].text = "Найденная формула (фрагмент протокола)"
    for i, (fn, f) in enumerate(forms, start=1):
        table2.rows[i].cells[0].text = fn
        table2.rows[i].cells[1].text = f if f else "—"

    # pass rate
    txt = result_path.read_text(encoding="utf-8")
    m = re.search(r"Benchmark pass rate:\s*(\d+)/(\d+)", txt)
    if m:
        add_paragraph(doc, f"Итог контрольной сводки: {m.group(1)} из {m.group(2)} наборов удовлетворяют заданным критериям.")


def append_full_protocol(doc: Document, result_path: Path) -> None:
    doc.add_page_break()
    add_heading(doc, "Приложение Д. Полный текст протокола result.txt", 1)
    add_paragraph(
        doc,
        "Ниже приведён текст файла протокола без сокращения (для архивации результатов вычислительного эксперимента).",
    )
    if not result_path.exists():
        return
    p = doc.add_paragraph()
    run = p.add_run(result_path.read_text(encoding="utf-8"))
    run.font.name = "Courier New"
    run.font.size = Pt(10)
    p.paragraph_format.first_line_indent = Cm(0)


def main() -> int:
    doc = Document()
    set_doc_defaults(doc)

    add_title_page_block(doc, DIPLOM / "01_titul_referat.md")
    add_referat(doc, DIPLOM / "01_titul_referat.md")
    doc.add_page_break()

    rp = RESULT_TXT if RESULT_TXT.exists() else Path(sys.argv[1]) if len(sys.argv) > 1 else RESULT_TXT
    ensure_fig_4_2_png(rp)

    sequence = [
        "02_soderzhanie.md",
        "03_vvedenie.md",
        "glava_1.md",
        "glava_2.md",
        "glava_3.md",
        "glava_4.md",
    ]
    for name in sequence:
        process_markdown_file(doc, DIPLOM / name)
        doc.add_page_break()

    add_experiment_section(doc, rp)
    doc.add_page_break()

    process_markdown_file(doc, DIPLOM / "zaklyuchenie.md")
    doc.add_page_break()

    add_heading(doc, "Примечание об оформлении (методические ориентиры)", 1)
    process_markdown_file(doc, DIPLOM / "00_oformlenie_gost.md")
    doc.add_page_break()

    process_markdown_file(doc, DIPLOM / "literatura.md")
    doc.add_page_break()

    process_markdown_file(doc, DIPLOM / "prilozheniya.md")
    doc.add_page_break()

    append_full_protocol(doc, rp)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"Saved: {OUT}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
